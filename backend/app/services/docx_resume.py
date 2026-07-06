from docx import Document

import tempfile


def create_docx(data):

    doc = Document()

    personal = data["personal"]

    doc.add_heading(

        personal["fullName"],

        level=1

    )

    doc.add_paragraph(

        personal["email"]

    )

    doc.add_paragraph(

        personal["phone"]

    )

    doc.add_heading(

        "Summary",

        level=2

    )

    doc.add_paragraph(

        personal["summary"]

    )

    doc.add_heading(

        "Skills",

        level=2

    )

    for skill in data["skills"]:

        doc.add_paragraph(

            skill,

            style="List Bullet"

        )

    doc.add_heading(

        "Experience",

        level=2

    )

    for exp in data["experience"]:

        doc.add_heading(

            exp["role"],

            level=3

        )

        doc.add_paragraph(

            exp["company"]

        )

        doc.add_paragraph(

            exp["description"]

        )

    file = tempfile.NamedTemporaryFile(

        delete=False,

        suffix=".docx",

    )

    doc.save(file.name)

    return file.name